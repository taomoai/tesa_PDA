import logging
import os
import re
import subprocess
import tempfile
from typing import Dict, Any, List

import html2text
import mammoth
from bs4 import BeautifulSoup
from docx import Document

from fastapi_app.utils.parsers.table_processor import TableProcessor

logger = logging.getLogger(__name__)

class WordFileParser():
    """Parser for Word files (DOC, DOCX)."""

    def __init__(self):
        """初始化 Word 文件解析器"""
        self.table_processor = TableProcessor()

    def _extract_headers_footers(self, file_path: str) -> Dict[str, str]:
        """
        使用python-docx提取Word文档的页眉页脚内容

        Args:
            file_path: Word文档路径

        Returns:
            包含页眉页脚内容的字典
        """
        headers_footers = {
            "headers": "",
            "footers": ""
        }

        try:
            doc = Document(file_path)

            # 提取所有节的页眉
            header_texts = []
            for section in doc.sections:
                header = section.header
                if not header.is_linked_to_previous:
                    # 提取页眉中的所有段落文本
                    for paragraph in header.paragraphs:
                        if paragraph.text.strip():
                            header_texts.append(paragraph.text.strip())

            # 合并页眉页脚文本
            headers_footers["headers"] = "\n".join(header_texts) if header_texts else ""
            logger.info(f"提取到页眉内容: {headers_footers['headers']}")

        except Exception as e:
            logger.warning(f"提取页眉页脚失败: {str(e)}")

        return headers_footers

    def parse(self, file_path: str) -> Dict[str, Any]:
        """
        Parse Word file and extract content.

        Args:
            file_path: The local path of the Word file to parse
            product: Optional OqcProduct instance containing product information

        Returns:
            Dict containing the extracted content with page_count
        """
        logger.info(f"开始处理Word文件: {file_path}")

        # 检查文件类型
        is_docx = file_path.lower().endswith('.docx')
        is_doc = file_path.lower().endswith('.doc')

        # 初始化一个基本的内容结构
        extracted_content = {"text": "", "tables": []}
        temp_path = None
        converted_path = None
        page_count = 0

        try:
            # 下载文件到临时目录
            # 如果是.doc文件，尝试转换为.docx
            if is_doc:
                try:
                    converted_path = self._convert_doc_to_docx(file_path)
                except Exception as convert_error:
                    logger.error(f"转换.doc到.docx失败: {str(convert_error)}")
                    return {"error": f"转换.doc到.docx失败: {str(convert_error)}"}

                if converted_path:
                    logger.info(f".doc文件已转换为.docx: {converted_path}")
                    temp_path = converted_path
                    is_docx = True
            else:
                # 对于直接的.docx文件，使用原始文件路径
                temp_path = file_path

            # 使用python-docx处理.docx文件
            if is_docx and temp_path:
                try:
                    # 检查页数限制（Word 最多 3 页）
                    doc_for_check = Document(temp_path)
                    page_breaks = 0
                    for paragraph in doc_for_check.paragraphs:
                        if paragraph.paragraph_format.page_break_before:
                            page_breaks += 1

                    # 如果没有分页符，至少是 1 页
                    page_count = page_breaks + 1 if page_breaks > 0 else 1

                    if page_count > 3:
                        error_msg = f"Word 文档页数超过限制。当前页数: {page_count}, 最大允许页数: 3"
                        logger.error(error_msg)
                        return {
                            "type": "docx",
                            "error": error_msg,
                            "page_count": page_count,
                            "max_pages": 3,
                            "is_page_limit_exceeded": True
                        }

                    logger.info(f"Word 文档页数检查通过: {page_count} 页")

                    # 首先提取页眉页脚内容
                    headers_footers = self._extract_headers_footers(temp_path)

                    # 使用mammoth将DOCX转换为HTML
                    with open(temp_path, "rb") as docx_file:
                        # 步骤1: 使用Mammoth将DOCX转换为HTML
                        result = mammoth.convert_to_html(docx_file)
                        html_content = result.value

                        # 步骤2: 使用html2text将HTML转换为Markdown
                        h = html2text.HTML2Text()
                        # 基础配置
                        h.body_width = 0  # 禁用自动换行
                        h.unicode_snob = True  # 保持 Unicode 字符
                        h.escape_snob = True  # 不转义特殊字符
                        
                        # 表格相关配置
                        h.bypass_tables = False  # 不跳过表格处理
                        h.ignore_tables = False  # 不忽略表格
                        h.pad_tables = True  # 在单元格中添加空格填充
                        h.wrap_tables = True  # 在表格前后添加空行
                        h.skip_internal_links = True  # 跳过内部链接
                        h.emphasis_mark = '*'  # 使用 * 作为强调标记
                        h.strong_mark = '**'  # 使用 ** 作为加粗标记
                        h.ul_item_mark = '-'  # 使用 - 作为无序列表标记
                        
                        # 提取rowspan信息
                        rowspan_info = []
                        try:
                            soup = BeautifulSoup(html_content, 'html.parser')
                            for td in soup.find_all('td', rowspan=True):
                                rowspan = td['rowspan']
                                content = td.get_text(strip=True)
                                rowspan_info.append({
                                    'content': content,
                                    'rowspan': rowspan
                                })
                            logger.info(f"Found rowspan cells: {rowspan_info}")
                        except Exception as e:
                            logger.warning(f"Error extracting rowspan info: {str(e)}")
                        
                        # 链接和图片配置
                        h.protect_links = True  # 保护链接中的特殊字符
                        h.inline_links = True  # 使用内联链接格式
                        h.images_to_alt = True  # 使用图片的alt文本
                        
                        # 文本格式配置
                        h.single_line_break = True  # 保持单个换行
                        h.ignore_emphasis = False  # 保留强调标记
                        h.backquote_code_style = True  # 使用反引号代码样式

                        original_markdown = h.handle(html_content)
                        logger.info(f"html_content: {html_content}")
                        logger.info(f"original_markdown: {original_markdown}")
                        
                        # 处理HTML内容以在表格单元格中的段落之间添加换行符
                        # 这可以解决表格单元格中相邻<p>标签被合并的问题
                        
                        # 替换表格单元格中的多个段落，使用<br/>标签连接内容
                        def replace_paragraphs_in_cell(match):
                            cell_content = match.group(1) if match.group(1) else ""
                            
                            # 如果单元格内容为空，直接返回原始匹配
                            if not cell_content.strip():
                                return match.group(0)
                            
                            # 如果内容中有换行标记，保持换行格式
                            if "<br/>" in cell_content or "<br>" in cell_content:
                                cell_content = cell_content.replace("<br>", "<br/>")
                                return f"<td>{cell_content}</td>"
                            
                            # 处理段落标记，将多个p标签内容用空格连接
                            if "<p>" in cell_content and "</p>" in cell_content:
                                # 提取所有段落内容
                                paragraphs = re.findall(r'<p>(.*?)</p>', cell_content, re.DOTALL)
                                # 用空格连接段落内容
                                new_content = " ".join(p.strip() for p in paragraphs if p.strip())
                                return f"<td>{new_content}</td>"
                            
                            # 如果没有特殊标记，返回原始内容
                            return match.group(0)
                        
                        # 对所有表格单元格应用替换
                        # 修改正则表达式以匹配带有属性的td标签
                        html_content = re.sub(r'<td(?:\s+[^>]*)?>(.*?)</td>', replace_paragraphs_in_cell, html_content, flags=re.DOTALL)

                        markdown = h.handle(html_content)
                        
                        # 在提取文本之前应用跨行处理
                        logger.info(f"处理前: {markdown}")
                        if rowspan_info:
                            logger.info(f"开始处理跨行信息: {rowspan_info}")
                            for rowspan in rowspan_info:
                                logger.info(f"rowspan: {rowspan}")
                                # 标准化内容格式：移除括号前后的空格
                                if '(' in rowspan['content'] and ')' in rowspan['content']:
                                    # 先标准化 rowspan 的内容
                                    normalized_content = rowspan['content'].replace(' (', '(').replace('( ', '(')
                                    normalized_content = normalized_content.replace(' )', ')').replace(') ', ')')
                                    rowspan['content'] = normalized_content
                                    # 转义括号
                                    rowspan['content'] = rowspan['content'].replace('(', '\\(').replace(')', '\\)')
                                    logger.info(f"标准化和转义后的rowspan: {rowspan}")
                                markdown = self.apply_rowspans_to_markdown(markdown, [rowspan])

                            logger.info(f"处理后的markdown: {markdown}")

                        # 提取和处理表格
                        tables = self.table_processor.extract_all_tables(markdown)
                        logger.info(f"提取到 {len(tables)} 个表格")

                        # 为 AI 格式化表格
                        formatted_tables = []
                        for i, table in enumerate(tables):
                            # 验证表格
                            is_valid, errors = self.table_processor.validate_table(table)
                            if not is_valid:
                                logger.warning(f"表格 {i+1} 验证失败: {errors}")

                            # 转换为 JSON 格式
                            json_table = self.table_processor.convert_to_json(table)
                            formatted_tables.append(json_table)

                            logger.info(f"表格 {i+1}: 类型={table.get('type')}, 行数={len(table.get('rows', []))}, 列数={len(table.get('headers', []))}")

                        # 合并页眉页脚内容到主体内容中
                        combined_text = ""

                        # 添加页眉内容（如果存在）
                        if headers_footers["headers"]:
                            combined_text += "=== 页眉内容 ===\n"
                            combined_text += headers_footers["headers"] + "\n\n"

                        # 添加主体内容
                        combined_text += "=== 文档内容 ===\n"
                        combined_text += markdown

                        # 返回提取的内容
                        return {
                            "type": "docx",
                            "text": combined_text,
                            "rowspans": rowspan_info,
                            "tables": formatted_tables,
                            "table_count": len(formatted_tables),
                            "page_count": page_count,
                            "is_page_limit_exceeded": False
                        }
                    
                except Exception as docx_error:
                    logger.error(f"使用python-docx处理文档失败: {str(docx_error)}", exc_info=True)
                    return {"error": f"处理Word文档失败: {str(docx_error)}"}
            return {"error": "无法处理文件"}
        
        except Exception as word_error:
            logger.error(f"Word处理错误: {str(word_error)}", exc_info=True)
            return {"error": f"Word文件处理失败: {str(word_error)}"}
        
        finally:
            # 清理临时文件
            if temp_path and os.path.exists(temp_path):
                try:
                    os.unlink(temp_path)
                    logger.info(f"已清理临时文件: {temp_path}")
                except Exception as e:
                    logger.warning(f"清理临时文件失败: {str(e)}")
            
            # 如果存在转换后的文件，也清理它
            if converted_path and os.path.exists(converted_path):
                try:
                    os.unlink(converted_path)
                    logger.info(f"已清理转换后的临时文件: {converted_path}")
                except Exception as e:
                    logger.warning(f"清理转换后的临时文件失败: {str(e)}")

    def _convert_doc_to_docx(self, doc_path: str) -> str:
        """
        将.doc文件转换为.docx文件。
        
        Args:
            doc_path: .doc文件的路径
            
        Returns:
            转换后的.docx文件路径
        """
        try:
            # 创建临时目录用于存放转换后的文件
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
                docx_path = temp_file.name
            
            # 使用 soffice 进行转换
            # 注意：这需要系统安装了 LibreOffice
            cmd = [
                'soffice',
                '--headless',
                '--convert-to',
                'docx',
                '--outdir',
                os.path.dirname(docx_path),
                doc_path
            ]
            
            # 执行转换命令
            process = subprocess.Popen(
                cmd,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE
            )
            stdout, stderr = process.communicate()
            
            if process.returncode != 0:
                raise Exception(f"转换失败: {stderr.decode()}")
            
            # 获取转换后的文件路径
            converted_path = os.path.join(
                os.path.dirname(docx_path),
                os.path.splitext(os.path.basename(doc_path))[0] + '.docx'
            )
            
            # 如果转换后的文件存在，返回其路径
            if os.path.exists(converted_path):
                return converted_path
            else:
                raise Exception("转换后的文件不存在")
            
        except Exception as e:
            logger.error(f"转换.doc到.docx失败: {str(e)}")
            raise 

    def apply_rowspans_to_markdown(self, markdown: str, rowspans: List[Dict[str, Any]]) -> str:
        """
        处理markdown表格中的rowspan信息，将跨行的内容复制到对应的空单元格中
        """
        if not rowspans:
            return markdown

        def normalize_for_comparison(text: str) -> str:
            """标准化文本以进行比较"""
            if not isinstance(text, str):
                text = str(text)
            # 移除所有转义字符和括号前后的空格
            text = text.replace('\\(', '(').replace('\\)', ')')
            text = text.replace(' (', '(').replace('( ', '(')
            text = text.replace(' )', ')').replace(') ', ')')
            text = text.strip()
            return text

        def format_content(content: str, line: str, source_line: str) -> str:
            """根据源行的格式处理内容"""
            if not isinstance(content, str):
                content = str(content)
            
            # 如果源行包含转义字符，我们也应该添加转义字符
            if '\\(' in source_line:
                # 先移除可能存在的转义字符
                content = content.replace('\\(', '(').replace('\\)', ')')
                # 然后添加转义字符
                content = content.replace('(', '\\(').replace(')', '\\)')
            
            # 获取源单元格的格式（空格）
            source_cells = source_line.split('|')
            if len(source_cells) > 2:
                source_cell = source_cells[-2]
                leading_spaces = len(source_cell) - len(source_cell.lstrip())
                trailing_spaces = len(source_cell) - len(source_cell.rstrip())
                content = ' ' * leading_spaces + content + ' ' * trailing_spaces
            
            return content

        lines = markdown.split('\n')
        
        # 对于每个跨行信息
        for span in rowspans:
            content = span['content']
            rowspan = int(span['rowspan'])
            
            # 找到内容所在的行
            start_line_idx = -1
            source_cell = None
            source_line = None
            
            # 查找源行（需要处理转义字符的情况）
            normalized_content = normalize_for_comparison(content)
            logger.info(f"Looking for normalized content: {normalized_content}")
            for i, line in enumerate(lines):
                cells = line.split('|')
                if len(cells) > 2:
                    cell_content = normalize_for_comparison(cells[-2])
                    logger.info(f"Comparing with cell content at line {i}: {cell_content}")
                    if normalized_content in cell_content:
                        start_line_idx = i
                        source_cell = cells[-2]
                        source_line = line
                        logger.info(f"Found match at line {i}: {source_cell}")
                        break
            
            if start_line_idx == -1:
                continue
            
            # 直接处理下一行
            next_line_idx = start_line_idx + 1
            logger.info(f"Checking next line {next_line_idx}")
            
            # 找到下一个完整的数据行
            while next_line_idx < len(lines):
                line = lines[next_line_idx]
                logger.info(f"Next line content: {line}")
                cells = line.split('|')
                logger.info(f"Cells in next line: {cells}")
                
                # 检查是否是完整的数据行
                if len(cells) >= 4:
                    logger.info(f"Last cell content: '{cells[-1]}'")
                    if not cells[-1].strip():  # 检查最后一个单元格是否为空
                        # 格式化内容
                        formatted_content = format_content(content, line, source_line)
                        logger.info(f"Formatted content: {formatted_content}")
                        
                        # 更新最后一列的内容
                        cells[-1] = formatted_content
                        logger.info(f"Updated cell: {cells[-1]}")
                        
                        # 重新组合行
                        new_line = '|'.join(cells)
                        lines[next_line_idx] = new_line
                        logger.info(f"Updated line {next_line_idx}: {new_line}")
                        break
                    else:
                        logger.info("Last cell is not empty, skipping")
                        break  # 如果找到完整行但最后一列不为空，也停止搜索
                else:
                    logger.info(f"Not enough cells in line: {len(cells)}, continuing to next line")
                    next_line_idx += 1  # 继续查找下一行
        return '\n'.join(lines)
